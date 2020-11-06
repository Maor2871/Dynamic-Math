import os
from random import choice, shuffle
from docx import Document
from docx.shared import Pt
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email import encoders
import smtplib


##### ---------- Functions ---------- #####

def user_input():
	"""
		The function receives and validates all the inputs that the user enters.
	"""
	
	# Ask for a minimum number.
	range_min = int(input("Please enter a minimum number: "))
	
	# Make sure that the received input is a number.
	if not range_min.isdigit():
		print("Next time enter a number please.")

	# Ask for a maximum number.
	range_max = int(input("Please enter a maximum number: "))
	
	# Make sure that the received input is a number.
	if not range_max.isdigit():
		print("Next time enter a number please.")
	
	# Ask for exercises count.
	amount = int(input("Please indicate the number of exercises you wish to generate: "))
	
	# Make sure that the received input is a number > 0.
	if not amount.isdigit() or amount.isdigit() and not amount > 0:
		print("Next time enter a number greater than zero please.")	
	
	# Ask for operations type.
	operations = input("Please indicate what operations you would like to have in the exercises.\nm-multiplication, d-division.\nYou can specify more than one operation, for example: dm (both ':' and 'x').\nOperations: ")
	
	# Make sure all the received operations are available.
	for chr in operations:	
		if chr not in operations_dict:
			print("The operation: " + chr + " is not available.\nPlease read again the instructions.")
	
	# Ask for the document name.
	file_name = input("worksheet name: ")
	
	# If the file name has a dot cut the name until the dot.
	if '.' in file_name: file_name = file_name[:file_name.index('.')]

	# The path of the new word document.
	file_path = os.path.join("Exercises", file_name + '.docx')
	
	# If the file name already exist, replace it with the new one.
	if os.path.exist(file_path):
		os.remove(file_path)
	
	# Return all the inputs.
	return range_min, range_max, amount, operations, file_name, file_path

def generate_riddles(amount, operation, range_min, range_max):
	"""
		The function receives operation type and amount of riddles to create and returns a list of operations.
	"""
	
	# Will contain all the exercises.
	riddles = []
	
	# Iterate over the amount of exercises expected.
	for i in range(amount):
	
		riddles.append(generate_riddle(operation, range_min, range_max))
	
	return riddles

def generate_riddle(operation, range_min, range_max):
	"""
		The function receives operation type and returns a riddle of that type.
	"""
	
	# Generate the numbers.
	num1 = choice(range(range_min, range_max))
	num2 = choice(range(range_min + 1, range_max))
	
	# If else operation type. Return the riddle of the received operation.
	if operation == "m":
	
		return str(num1) + ' x ' + str(num2) + '  = '
		
	elif operation == "d":
	
		return str(num1 * num2) + ' : ' + str(num1) + '  = '
		

def send_mail(mail_address, file_name, file_path):
	"""
		The function sends the document by mail.
	"""

	fromaddr = 'classifiedanonymouse@gmail.com'
	toaddr = mail_address


	msg = MIMEMultipart()
	msg['From'] = fromaddr
	msg['To'] = toaddr
	msg['Subject'] = file_name
	msg.attach(MIMEText('Dynamic math worksheet created ' + file_name + " for you.\nIt is attached to the mail as a word document."))


	attachment = open(file_path, 'rb')
	part = MIMEBase("application", "octet-stream")
	part.set_payload(attachment.read())
	encoders.encode_base64(part)
	part.add_header("Content-Disposition", "attachment; filename= " + file_name + ".docx")
	msg.attach(part)
	msg = msg.as_string()

	try:
		server = smtplib.SMTP('smtp.gmail.com:587')
		server.ehlo()
		server.starttls()
		server.login(fromaddr, 'ThisIsClassified4@')
		server.sendmail(fromaddr, toaddr, msg)
		server.quit()
		
		print('Email sent successfully')

	except Exception as e:
		print("Email couldn't be sent")
	

##### ---------- Create exercises ---------- #####


# ----- inputs -----

# Get all the necessary inputs from the user.
# range_min, range_max, amount, operations, file_name, file_path = user_input()

range_min = 2
range_max = 12
amount = 100
operations = 'md'
file_name = 'ws1'
file_path = os.path.join("Exercises", file_name + '.docx')
if os.path.exists(file_path): os.remove(file_path)

# Will contain all the exercises.
riddles = []

# For the header in the document.
operations_dict = {'m': "Multiplication", 'd': "Division"}

# Check if more than one operation selected.
if len(operations) > 1:
	amount = int(round(amount / len(operations)))

# Check if the operation of the exercises is division.
if "d" in operations:
	
	riddles += generate_riddles(amount, 'd', range_min, range_max)

# Create multiplication exercises as default.
if 'm' in operations:
	
	riddles += generate_riddles(amount, 'm', range_min, range_max)

shuffle(riddles)
		
#### ---------- Save and print ---------- #####

# Format the riddles.
riddles_formatted = '\n\n\n'.join(riddles[i] + '\t\t\t\t\t\t\t' + riddles[i - 1] for i in range(1, len(riddles), 2))

# If odd number of exercises add the last one.
if not len(riddles_formatted) %2 == 0: riddles_formatted += '\n\n\n' + riddles[-1]

# ----- Word print ----- #

# Create the word document.
document = Document()

# Add header to the document.
document.add_heading((', '.join(operations_dict[operation] for operation in operations)).replace(', ', " and ", -1) + ' of numbers from ' + str(range_min) + ' to ' + str(range_max) + '\n\n', 0)

# Define the font size.
style = document.styles['Normal']
font = style.font
font.size = Pt(14)

# Add the exercises to the document.
document.add_paragraph(riddles_formatted, style='Normal')

# Save the document.
document.save(file_path)

# Send the document to mail.
send_mail("maor29468@gmail.com", file_name, file_path)

# Print the word document.
#os.startfile(os.path.join("Exercises", file_name + '.docx'), "print")