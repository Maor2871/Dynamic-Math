import os
from sys import exit
from random import choice, shuffle
from docx import Document
from docx.shared import Pt
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email import encoders
import smtplib
from kivy.app import App
from kivy.uix.widget import Widget
from kivy.uix.gridlayout import GridLayout
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.textinput import TextInput
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from kivy.uix.checkbox import CheckBox
from kivy.core.window import Window
from kivy.properties import NumericProperty
from kivy.properties import ListProperty


class Tool():
	"""
		The class implements the tool functions.
	"""
	
	def __init__(self, operations, amount, range_min, range_max, file_path, file_name, mail_address, action):
		"""
		"""
		
		# Save all the necessary configurations.
		self.operations = operations
		self.amount = amount
		self.range_min = range_min
		self.range_max = range_max
		self.file_path = file_path
		self.file_name = file_name
		self.mail_address = mail_address
		self.action = action
		
		# Contain all the exercises.
		self.riddles = []
		
		# All the exercises formatted as a string.
		self.riddles_formatted = ""

	def create_exercises(self):
		"""
			The function creates and randomizes the desired exercises.
		"""

		# Check if more than one operation selected.
		if len(self.operations) > 1:
			amount = int(round(self.amount / len(self.operations)))

		# Check if division is a required operation.
		if "d" in operations:
			
			# Add to riddles division exercises.
			self.riddles += self.generate_riddles('d')

		# Check if multiplication is a required operation.
		if 'm' in operations:
			
			# Add to riddles multiplication exercises.
			self.riddles += generate_riddles(amount, 'm', range_min, range_max)
		
		# Shuffle all the exercises.
		shuffle(self.riddles)
	
	def generate_riddles(self, operation):
		"""
			The function receives operation type and amount of riddles to create and returns a list of operations.
		"""
		
		# All the riddles of the current operation.
		operation_riddles = []
		
		# Iterate over the amount of exercises expected.
		for i in range(self.amount):
			
			# Generate a single riddle of the form of the current operation and add it to the list.
			operation_riddles.append(self.generate_riddle(operation))
			
		# Return all the riddles of the received operation.
		return operation_riddles		

	def generate_riddle(self, operation):
		"""
			The function receives operation type and returns a riddle of that type.
		"""
		
		# Generate the numbers.
		num1 = choice(range(self.range_min, self.range_max))
		num2 = choice(range(self.range_min, self.range_max))
		
		# If else operation type. Return the riddle of the received operation.
		if operation == "m":
		
			return str(num1) + ' x ' + str(num2) + '  = '
			
		elif operation == "d":
		
			return str(num1 * num2) + ' : ' + str(num1) + '  = '

	def format_riddles(self):
		"""
			The function receives exercises list and formats them to a string.
		"""
		
		# Format the riddles.
		self.riddles_formatted = '\n\n\n'.join(self.riddles[i] + '\t\t\t\t\t\t\t' + self.riddles[i - 1] for i in range(1, len(self.riddles), 2))

		# If odd number of exercises add the last one.
		if not len(self.riddles_formatted) %2 == 0: self.riddles_formatted += '\n\n\n' + self.riddles[-1]	

	def save_as_word(self):
		"""
			The function saves the exercises as a word document.
		"""
		
		# Create the word document.
		document = Document()

		# Add header to the document.
		document.add_heading((', '.join(operations_dict[operation] for operation in self.operations)).replace(', ', " and ", -1) + ' of numbers from ' + str(self.range_min) + ' to ' + str(self.range_max) + '\n\n', 0)

		# Define the font size.
		style = document.styles['Normal']
		font = style.font
		font.size = Pt(14)

		# Add the exercises to the document.
		document.add_paragraph(self.riddles_formatted, style='Normal')

		# Save the document.
		document.save(self.file_path)

	def send_mail(self):
		"""
			The function sends the document by mail.
		"""
		
		fromaddr = 'classifiedanonymouse@gmail.com'
		toaddr = self.mail_address

		# Create the message
		msg = MIMEMultipart()
		msg['From'] = fromaddr
		msg['To'] = toaddr
		msg['Subject'] = self.file_name
		msg.attach(MIMEText('Dynamic math worksheet created ' + self.file_name + " for you.\nIt is attached to the mail as a word document."))
		
		# Attach the word document.
		attachment = open(self.file_path, 'rb')
		part = MIMEBase("application", "octet-stream")
		part.set_payload(attachment.read())
		encoders.encode_base64(part)
		part.add_header("Content-Disposition", "attachment; filename= " + self.file_name + ".docx")
		msg.attach(part)
		msg = msg.as_string()

		try:
			server = smtplib.SMTP('smtp.gmail.com:587')
			server.ehlo()
			server.starttls()
			server.login(fromaddr, 'ThisIsClassified4@')
			# server.sendmail(fromaddr, toaddr, msg)
			server.quit()
			
			return 'Email sent successfully.'

		except Exception as e:
			return "Email couldn't be sent."
			
	def execute_actions(self):
		"""
			receive the word document and choose what to do with it.
		"""
		
		actions_status = {"print": "", "mail": ""}
		
		# If the user want the document by mail send it to his mail address.
		if 'm' in action:

			# Send the document to mail and update the status to success or failure.
			actions_feedbacks["mail"] = send_mail(mail_address, file_name, file_path)

		# If the user want the document to be printed, send it to the printer.
		if 'p' in action:
			
			try:
			
				# Print the word document.
				#os.startfile(os.path.join("Exercises", file_name + '.docx'), "print")

				# Update the printing status to success.
				actions_feedbacks["print"] = "Document Sent to printer successfully."
				
			except Exception as e:
				
				# Update the printing status to failure.
				actions_feedbacks["print"] = "Document couldn't be sent to printer."
				
		# Return the results of the actions.
		return actions_status

	def abort(self):
		"""
			The function aborts the script.
		"""
		
		exit()


class LUI():
	"""
		The class runs the tool with line user interface.
	"""
	
	def __init__(self):
	
		# Indicates what operations are available and their short-cuts.
		self.operations_dict = {'m': "Multiplication", 'd': "Division"}

		# Get all the necessary inputs from the user and initialize the tool.
		self.tool = Tool(self.user_input())
		
		# Start the tool.
		self.start_tool()
		
	def start_tool(self):
		"""
			The function starts the tool with the user configurations.
		"""
		
		print("Creating exercises... ", end="")

		# Create all the exercises.
		self.toolcreate_exercises()

		print("complete.\n")
		print("formatting exercises... ", end="")

		# Format all the exercises to a string.
		self.tool.format_riddles()

		print("complete.\n")
		print("\nSaving exercises to a word document... ", end="")

		# Save the exercises as a word document.
		self.tool.save_as_word(self)

		print("complete.\n")
		print("Executing action...")

		# Execute the actions that the user wishes to do with the document.
		self.actions_feedbacks(self.tool.execute_actions())

		print("\nAll actions are done successfully, thank you for using Dynamic Math.\nCome again soon.")
		
	def actions_feedbacks(self, actions_feedbacks):
		"""
			The function receives indications if the actions were completed successfully and sends a feedback to the user.
		"""
		
		# Print the results of the executed actions.
		for action in actions_feedbacks:
			
			# If the action was executed.
			if actions_feedbacks[action]:
				
				# Print its results.
				print(actions_feedbacks[action])

	def user_input(self):
		"""
			The function receives and validates all the inputs that the user enters.
		"""
		
		# Ask for a minimum number.
		range_min = input("Please enter a minimum number: ")
		
		# Make sure that the received input is a number.
		if not range_min.isdigit():
			print("Next time enter a number please.")
			abort()

		# Ask for a maximum number.
		range_max = input("\nPlease enter a maximum number: ")
		
		# Make sure that the received input is a number.
		if not range_max.isdigit():
			print("Next time enter a number please.")
			abort()
		
		# Convert range min and range max to int numbers.
		range_min = int(range_min)
		range_max = int(range_max)
		
		# Ask for exercises count.
		amount = input("\nPlease indicate the number of exercises you wish to generate: ")
		
		# Make sure that the received input is a number > 0.
		if not amount.isdigit() or amount.isdigit() and not int(amount) > 0:
			print("Next time enter a number greater than zero please.")
			abort()
			
		# Convert amount to int number.
		amount = int(amount)
		
		# Ask for operations type.
		operations = input("\nPlease indicate what operations you would like to have in the exercises.\n\nm-multiplication\nd-division.\n\n- You can specify more than one operation, for example: dm (both ':' and 'x').\n\nOperations: ")
		
		# Make sure all the received operations are available.
		for chr in operations:	
			if chr not in operations_dict:
				print("The operation: " + chr + " is not available.\nPlease read again the instructions.")
				abort()
		
		# Ask for the document name.
		file_name = input("\nworksheet name: ")
		
		# If the file name has a dot cut the name until the dot.
		if '.' in file_name: file_name = file_name[:file_name.index('.')]

		# The path of the new word document.
		file_path = os.path.join("Exercises", file_name + '.docx')
		
		# If the file name already exist, replace it with the new one.
		if os.path.exists(file_path):
			os.remove(file_path)
		
		# Ask for the desired output: p for print or m for mail.
		action = input("\nPlease indicate in what form you want to get the file.\n\np- direct print\nm- Send document to mail.\n\n- You can specify more than one action, for example: mp (both printing and mailing).\n\nActions: ")
		
		# Make sure an available action received.
		if not 'p' in action and not 'm' in action:
			print("You must enter either 'p' or 'm' to indicate if you want the file to be sent to the printer or to your mail address.")
			abort()
		
		# If the user wants the file to be sent by mail, ask him for his mail address.
		if 'm' in action:
			
			# Get the mail address of the user.
			mail_address = input("\nPlease enter your mail address: ")
			
			# Make a minimal check that a valid mail address was entered.
			if not '@' in mail_address:
				print("A mail address must have '@' in it.")
				abort()
		
		# If not set mail address to None.
		else:
		
			mail_address = None
		
		# Indicate the user that the all the input was received.
		print('\n\nAll input received successfully.\n')
		
		# Return all the inputs.
		return operations, amount, range_min, range_max, file_path, file_name, mail_address, action


class GUI(App):
	"""
		The class runs the tool with graphical user interface.
	"""

	def __init__(self):
		"""
			Initialize the GUI interface.
		"""

		super(GUI, self).__init__()

		# The main layout of the GUI.
		self.main_layout = GUI.MainLayout()
		
	def build(self):
		"""
			Returns the main layout.
		"""
			
		# Set the GUI to full-screen.
		#Window.fullscreen = True
		Window.size = (1080, 1920)

		return self.main_layout
		
	class MainLayout(BoxLayout):
		"""
			The class represents the main layout of the GUI.
		"""
		
		def __init__(self):
			"""
				Initialize the main layout.
			"""
			
			# Initialize the box layout.
			super(GUI.MainLayout, self).__init__(orientation="vertical")
			
			# The exercises properties panel.
			self.exercises_properties_panel = GUI.MainLayout.ExercisesPropertiesPanel()
			
			# Get the mail address of the user.
			self.mail_address = GUI.Abstracts.Input(title="Please enter your mail address", orientation="vertical", default_text="Mail address here")
			
			# Ask for the desired name of the document.
			self.document_name = GUI.Abstracts.Input(title="Please enter desired document name", orientation="vertical", default_text="Document name here")
			
			# Arrange the mail address and the document name layout.
			self.document_mail_layout = BoxLayout(orientation="vertical")
			self.document_mail_layout.add_widget(self.exercises_properties_panel)
			self.document_mail_layout.add_widget(self.mail_address)
			self.document_mail_layout.add_widget(self.document_name)
			self.bottom_layout = BoxLayout(orientation="horizontal")
			self.bottom_layout.add_widget(BoxLayout())
			self.bottom_layout.add_widget(self.document_mail_layout)
			self.bottom_layout.add_widget(BoxLayout())
			
			# Add all the layouts to the main layout.
			#self.add_widget(self.exercises_properties_panel)
			self.add_widget(self.bottom_layout)
			
			self.mail_address.update_padding()
	
		class ExercisesPropertiesPanel(BoxLayout):
			"""
				The class represents a panel which allows the user to define the properties of the exercises he wants in the document.
			"""
			
			def __init__(self):
				"""
					Initializes the exercises properties panel layout.
				"""
				
				# Initialize the box layout.
				super(GUI.MainLayout.ExercisesPropertiesPanel, self).__init__(orientation='vertical', size_hint=(1, 3))
				
				# --- Components --- #
				
				# The minimum number input.
				self.min_number_input = GUI.Abstracts.NumericInput(title="minimum number", min_value=-1000000000, max_value=1000000000, default_text='0')

				# The maximum number input.
				self.max_number_input = GUI.Abstracts.NumericInput(title="maximum number", min_value=-1000000000, max_value=1000000000, default_text='0')
				
				# The operations that the user wants in the exercises.
				self.operations_layout = GUI.Abstracts.CheckBoxes(title="Exercises Operations", checkboxes_titles=["Multiplication", "Division"], orientation="vertical", checkboxes_orientation="horizontal")
				
				# The amount of exercises the user wants to create.
				self.exercises_amount = GUI.Abstracts.NumericInput(title="How many exercises to create?", min_value=1, max_value=10000, orientation="vertical", default_text='1')
				
				# --- Layouts configurations --- #
				
				# A layout to the numbers range.
				self.range_layout = BoxLayout(orientation="horizontal")
				self.range_layout.add_widget(self.min_number_input)
				self.range_layout.add_widget(self.max_number_input)
				
				# Add the numbers range and the operations to the layout.
				self.add_widget(self.range_layout)
				self.add_widget(self.operations_layout)
				self.add_widget(self.exercises_amount)
				
			def get_current_values(self):
				"""
					The function returns the current values of all the exercises properties.
				"""
				
				return self.min_number_input.text, self.max_number_input.text, self.operations_layout.get_checked(), self.exercises_amount.text

	class Abstracts():
		"""
			The class contain generic layouts.
		"""
		
		class Input(BoxLayout):
			"""
				The class represents a numeric input layout. It has an input text box and a title.
			"""
			
			def __init__(self, title, default_text="Text Here", orientation='vertical', multiline=False, max_length=None):
				"""
					Initialize the input layout.
				"""
				
				# Call the box layout initializer.
				super(GUI.Abstracts.Input, self).__init__(orientation=orientation)
				
				# Create the title.
				self.title = Label(text=title, pos_hint={'center_x': .5})
				
				# The default text when text input in empty.
				self.default_text = default_text
				
				# Create the text input.
				self.text_input = TextInput(size_hint=(.3, 3), pos_hint={'center_x': .5})
				
				# Align the text in the text input to center.
				self.text_input.bind(text=self.update_padding)
				self.text_input.bind(size=self.update_padding)
				self.text_input.bind(focus=self.on_focus)
				#self.text_input.bind(touch_up=self.on_touch_up)
	
				# Set the input text to be white.
				self.text_input.background_color = (0, 0, 0, 0)
				self.text_input.foreground_color = (1, 1, 1, 1)
				
				# Initialize the text input to center.
				self.update_padding()
				
				# Add the widgets to the layout.
				self.add_widget(self.title)
				self.add_widget(self.text_input)
				
				self.text_input.hint_text = self.default_text

			def update_padding(self, instance=None, value=None):
				"""
					Align the text to center.
				"""
				
				if not self.text_input.text:

					text_input_text = self.default_text
				
				else:
				
					text_input_text = self.text_input.text
				
				# Calculate the text width.
				text_width = self.text_input._get_text_width(text_input_text, self.text_input.tab_width, self.text_input._label_cached)

				# Update the text position to point on the center.
				self.text_input.padding_x = (self.text_input.width - text_width)/2
			
			def on_focus(self, instance, value=None):
				"""
					The function is being called when the user clicks the text input.
				"""
				
				if self.text_input.text == self.default_text:

					self.text_input.text_hint=""
					self.text = ""
					
			def on_touch_up(self, instance, value=None):
				"""
					The function is being called the text input in unfocused.
				"""
				
				# If text input is empty set its text to default text.
				if not self.text_input.text:
					self.text_input.text_hint = self.default_text
				
		class NumericInput(Input):
			"""
				The class represents an input of a number.
			"""
			
			def __init__(self,title, min_value, max_value, orientation='vertical', default_text='0'):
				"""
					Initialize the numeric input layout.
				"""
				
				super(GUI.Abstracts.NumericInput, self).__init__(title=title, orientation=orientation, default_text=default_text)
				
				# Update the minimum value and the maximum value of the number.
				self.min_value = min_value
				self.max_value = max_value
				
				# Set the input type to a number.
				self.input_type = 'number'
				self.input_filter = 'int'

				# Validate the number entered.
				self.text_input.bind(text=self.validate_number)
			
			def validate_number(self, instance, numeric_input):
				"""
					The function is being called when the value changes.
				"""

				# Make sure that the numeric value is within the boundaries.
				if self.text_input.text and self.text_input.text.isdigit() and not self.min_value <= int(self.text_input.text) <= self.max_value:

					# Create a pop-up that tells the user he exceeded the boundaries.
					popup = Popup(title='Attention', content=Label(text="value must be between " + str(self.min_value) + " and " + str(self.max_value)), size_hint=(None, None), size=(300, 200))
					
					# Empty the value.
					self.text_input.text = self.default_text
					
					# Pop-up the pop-up.
					popup.open()
		
		class CheckBoxes(BoxLayout):
			"""
				The class represents a set of check-boxes.
			"""
			
			def __init__(self, title, checkboxes_titles, orientation="vertical", checkboxes_orientation="horizontal"):
				"""
					The function initializes the check-boxes class.
				"""
				
				# Initialize the box-layout.
				super(GUI.Abstracts.CheckBoxes, self).__init__(orientation=orientation)
				
				# All the check-boxes titles.
				self.checkboxes_titles = checkboxes_titles
				
				# Create all the check-boxes and store them in a list.
				self.checkboxes = [GUI.Abstracts.CheckBoxes.CheckBoxTitle(title=title) for title in checkboxes_titles]
				
				# The title of the check-boxes group.
				self.title = Label(text=title)
				
				# The layout of the check-boxes themselves.
				self.checkboxes_layout = BoxLayout(orientation=checkboxes_orientation)
				
				# Create and add all the check-boxes to their layout.
				[self.checkboxes_layout.add_widget(checkbox) for checkbox in self.checkboxes]
				
				# Add the title and the check-boxes to the layout.
				self.add_widget(self.title)
				self.add_widget(self.checkboxes_layout)
				
				
			class CheckBoxTitle(BoxLayout):
				"""
					The class represents a check-box with title.
				"""
				
				def __init__(self, title, orientation="horizontal"):
					"""
						The function initializes the check-box-title.
					"""
					
					# Initialize the box layout.
					super(GUI.Abstracts.CheckBoxes.CheckBoxTitle, self).__init__(orientation=self.orientation)

					# The orientation of the box layout. default set to horizontal.
					self.orientation = orientation
					
					# The check-box itself.
					self.checkbox = CheckBox()
					
					# The title of the check-box.
					self.title = Label(text=title)
					
					
					# Add the check-box and the title to the layout.
					self.add_widget(self.title)
					self.add_widget(self.checkbox)

# Start the game.
GUI().run()